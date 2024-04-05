import PyPDF2
from docx import Document
import re
from docx import Document
import json
import json
from rdflib import Graph, Literal, Namespace, RDF, URIRef
import uuid

## This code will convert pdf data to output.docx file format
def extract_text_from_pdf(pdf_path, start_page=13, end_page=None):
    text = ""
    with open(pdf_path, "rb") as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        if end_page is None or end_page > len(pdf_reader.pages):
            end_page = len(pdf_reader.pages)
        
        for page_number in range(start_page, end_page):
            page = pdf_reader.pages[page_number]
            # Replace bullet points with blank spaces
            page_text = page.extract_text().replace('•', '')
            text += page_text
    return text

# Function to write text to a Word document
def write_text_to_word(text, output_word_path):
    document = Document()
    document.add_paragraph(text)
    document.save(output_word_path)


# Example PDF file path (replace this with your actual PDF file path)
pdf_path = "C://Users//aksha//Downloads//Courses PDF's//Web/AB_31_2015Teil1 (1).pdf"

# Extract text from the PDF
pdf_text = extract_text_from_pdf(pdf_path)

# Example Word document path (replace this with your desired output path)
output_word_path = "output.docx"

# Write extracted text to Word document
write_text_to_word(pdf_text, output_word_path)

print(f"Text from PDF has been written to {output_word_path}")

## This code section will take output.docx file and convert it into JSON file
# Regular expressions to extract information
module_number_pattern = re.compile(r"  (\S+ [-\S]*)")
module_name_pattern = re.compile(r"Modulname\s*(.+)")
contents_pattern = re.compile(r"Inhalte\s*:\s*(.+?)\s*Qualifikationsziele\s*:", re.DOTALL)
credit_points_pattern = re.compile(r"Leistungspunkte\s*u?nd\s*Noten\D*(\d+)")
work_load_pattern = re.compile(r"Arbeitsaufwand\D*(\d+)\s*(?:\n\s*)*AS")

# Function to extract information from a page
def extract_information(page_text):
    result = {}

    # Extract Module Number
    match_module_number = module_number_pattern.search(page_text)
    if match_module_number:
        contents = match_module_number.group(1).strip()
        # Remove Unwanted Space
        contents = contents.replace(' ', '')
        result["Modulnummer"] = contents

    # Extract Module Name
    match_module_name = module_name_pattern.search(page_text)
    if match_module_name:
        result["Modulname"] = match_module_name.group(1)

    # Extract Contents
    match_contents = contents_pattern.search(page_text)
    if match_contents:
        contents = match_contents.group(1).strip()
        # Remove unwanted characters
        contents = contents.replace('\n', '')
        result["Inhalte und Qualifikationsziele"] = contents

    # Extract Credit Points
    match_credit_points = credit_points_pattern.search(page_text)
    if match_credit_points:
        result["Leistungspunkte und Noten"] = match_credit_points.group(1)

    # Extract Workload
    match_work_load = work_load_pattern.search(page_text)
    if match_work_load:
        result["Arbeitsaufwand"] = match_work_load.group(1)

    return result

# Function to read text from a Word document
def read_text_from_word(word_path):
    doc = Document(word_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

# Example Word document path (replace this with your actual Word file path)
input_word_path = "output.docx"

# Read text from the Word document
word_text = read_text_from_word(input_word_path)

# Split the text into pages
pages = re.split(r"Modulnummer", word_text)[1:]

# Extract information from each page
result_list = [extract_information(page) for page in pages]

# Convert the list of dictionaries to JSON
json_data = json.dumps(result_list, indent=2)

# Write JSON data to a separate file
output_json_path = "output.json"
with open(output_json_path, "w") as json_file:
    json_file.write(json_data)

print(f"JSON data has been written to {output_json_path}")

# Load JSON data
data = json.loads(json_data)

# RDF Namespace
module_ns = URIRef("http://tuc.web.engineering/module#")

# Create RDF graph
g = Graph()

# Iterate over each module in the JSON data
for module in data:
    module_number = module.get("Modulnummer", "")
    module_name = module.get("Modulname", "")
    content = module.get("Inhalte und Qualifikationsziele", "")
    credit_points = module.get("Leistungspunkte und Noten", "")
    work_load = module.get("Arbeitsaufwand", "")

    # Generate a UUID based on the current timestamp and node (hardware address)
    module_uuid = uuid.uuid1()

    # Convert the UUID to a string
    module_uuid_str = str(module_uuid)
    # RDF URI for the module
    module_uri = URIRef(f"http://tuc.web.engineering/module#{module_uuid_str}")

    # Add RDF triples for the module
    g.add((module_uri, RDF.type, module_ns))
    g.add((module_uri, URIRef("http://tuc.web.engineering/module#hasModuleNumber"), Literal(module_number, datatype="http://www.w3.org/2001/XMLSchema#string")))
    g.add((module_uri, URIRef("http://tuc.web.engineering/module#hasName"), Literal(module_name, datatype="http://www.w3.org/2001/XMLSchema#string")))
    g.add((module_uri, URIRef("http://tuc.web.engineering/module#hasContent"), Literal(content, datatype="http://www.w3.org/2001/XMLSchema#string")))
    
    # Check if credit_points is non-empty before converting to integer
    if credit_points:
        try:
            credit_points_value = int(credit_points)
            g.add((module_uri, URIRef("http://tuc.web.engineering/module#hasCreditPoints"), Literal(credit_points_value, datatype="http://www.w3.org/2001/XMLSchema#integer")))
        except ValueError as ve:
            print(f"Error converting credit_points to integer for module {module_name}: {ve}")

     # Check if work_load is non-empty before converting to integer
    if work_load:
        try:
            work_load_value = int(work_load)
            g.add((module_uri, URIRef("http://tuc.web.engineering/module#hasWorkLoad"), Literal(work_load_value, datatype="http://www.w3.org/2001/XMLSchema#integer")))
        except ValueError as ve:
            print(f"Error converting work_load to integer for module {module_name}: {ve}")

    # Add additional RDF triples for each module
    g.add((module_uri, URIRef("http://tuc/web/engineering/module#hasUniversity"), URIRef("http://across/university#TUC")))
    g.add((module_uri, URIRef("http://tuc/course#hasCourse"), URIRef("http://tuc/course#WebEngineering")))

# Serialize RDF graph to RDF/XML format
rdf_outputBytes = (g.serialize(format="xml")).encode('utf-8')

# Save RDF data to a file with proper encoding
output_rdf_path = "output.rdf"
with open(output_rdf_path, "wb") as rdf_file:
    rdf_file.write(rdf_outputBytes)

print(f"RDF data has been saved to {output_rdf_path}")
